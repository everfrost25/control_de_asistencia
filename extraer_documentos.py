from docx import Document

files = [
    'ACTA_CONTROL_ASISTENCIA.docx',
    'EDT_CONTROL_ASISTENCIA.docx',
    'CRONOGRAMA_CONTROL_ASISTENCIA.docx',
    'MODELO DE COSTO – SISTEMA WEB DE CONTROL DE ASISTENCIA.docx',
]

for filename in files:
    document = Document(filename)
    print(f'\n### {filename}')
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            print(paragraph.text)
    for index, table in enumerate(document.tables, start=1):
        print(f'TABLA {index}')
        for row in table.rows:
            print(' | '.join(cell.text.replace('\n', ' / ') for cell in row.cells))
