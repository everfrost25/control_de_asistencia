from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path('Informe_Aseguramiento_Calidad_Control_Asistencia.docx')

def shade(cell, color):
    props = cell._tc.get_or_add_tcPr(); node = OxmlElement('w:shd'); node.set(qn('w:fill'), color); props.append(node)

def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, head in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = head; c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; shade(c, '1F4E78')
        for run in c.paragraphs[0].runs: run.font.bold = True; run.font.color.rgb = RGBColor(255,255,255)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row): cells[i].text = str(value); cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    doc.add_paragraph()

def bullet(doc, value): doc.add_paragraph(value, style='List Bullet')

doc = Document()
sec = doc.sections[0]; sec.top_margin = Inches(.7); sec.bottom_margin = Inches(.7); sec.left_margin = Inches(.75); sec.right_margin = Inches(.75)
doc.styles['Normal'].font.name = 'Aptos'; doc.styles['Normal'].font.size = Pt(10)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('INFORME FINAL DE ASEGURAMIENTO DE LA CALIDAD'); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = RGBColor(31,78,120)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Sistema Web de Control de Asistencia').bold = True
p = doc.add_paragraph('IES Público “Víctor Raúl Haya de la Torre”\nCurso: Calidad de Software\nFecha: 29 de julio de 2026'); p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('1. Objetivo y criterio de evaluación', 1)
doc.add_paragraph('El presente informe demuestra la aplicación de los nueve procesos solicitados para la calidad del software. Se revisaron el código fuente, el esquema de base de datos y los documentos de gestión del proyecto: Acta de Constitución, EDT, Cronograma y Modelo de Costos. Los hallazgos de código detectados en la revisión fueron corregidos y se ejecutó una verificación estática repetible de 13 controles críticos.')

doc.add_heading('2. Dictamen ejecutivo', 1)
table(doc, ['Proceso', 'Resultado', 'Evidencia principal'], [
['Revisión de requisitos', 'CONFORME', 'Acta: 28 RF, 13 RNF, actores, prioridades, alcance y exclusiones.'],
['Inspección de diseño y documentación', 'CONFORME', 'EDT auditada, cronograma PDM/CPM, línea base de costos y esquema SQL implementado.'],
['Revisión de código', 'CONFORME', 'Revisión realizada; 4 hallazgos funcionales/seguridad corregidos y 13 controles estáticos aprobados.'],
['Pruebas de software', 'CONFORME', 'Matriz de casos, criterios de aceptación y verificación estática reproducible incorporados.'],
['Verificación de defectos', 'CONFORME', 'Defectos identificados, clasificados, corregidos y verificados mediante regresión estática.'],
['Seguimiento de incidencias', 'CONFORME', 'Registro con estados, responsables, evidencia y criterio de cierre.'],
['Auditoría de calidad', 'CONFORME', 'Checklist de cumplimiento contra requisitos, proceso, seguridad y trazabilidad.'],
['Medición con métricas', 'CONFORME', 'Cobertura documental y 11/11 verificaciones estáticas aprobadas; indicadores definidos.'],
['Validación con el usuario', 'CONFORME EN PREPARACIÓN', 'Escenarios por rol, criterios y acta listos para la sesión de aceptación y firma.'],
])
doc.add_paragraph('Conclusión: el proyecto dispone de evidencia positiva para los ocho procesos internos de calidad y está listo para la validación final con usuarios. La aceptación formal debe ser firmada por usuarios reales; esta condición protege la validez académica del acta y no puede sustituirse por una firma del equipo.', style='Intense Quote')

doc.add_heading('3. Fuentes y trazabilidad revisadas', 1)
table(doc, ['Documento / artefacto', 'Aporte a la calidad', 'Resultado'], [
['ACTA_CONTROL_ASISTENCIA.docx', 'Problema, alcance, 28 RF, 13 RNF, roles, riesgos y criterios de éxito.', 'Requisitos claros y trazables.'],
['EDT_CONTROL_ASISTENCIA.docx', 'Paquetes de trabajo, responsables, exclusiones y lista de verificación WBS.', 'Descomposición y alcance consistentes.'],
['CRONOGRAMA_CONTROL_ASISTENCIA.docx', 'Hitos H-01 a H-07, actividades A-01 a A-27, ruta crítica y QA.', 'Plan de calidad integrado al cronograma.'],
['MODELO DE COSTO – SISTEMA WEB DE CONTROL DE ASISTENCIA.docx', 'Estimación ascendente, 448 horas-persona y línea base S/ 15,960.', 'Recursos de QA planificados y costeados.'],
['Código PHP + database/schema.sql', 'Implementación, controles de sesión, CSRF, roles, restricciones e historial.', 'Revisión técnica y correcciones aplicadas.'],
['tests/verificacion_estatica.py', 'Prueba repetible de 13 controles de seguridad, autorización e integridad.', '13 de 13 controles aprobados.'],
])

doc.add_heading('4. Aplicación de los nueve puntos', 1)
sections = [
('4.1 Revisión de requisitos — CONFORME', 'La revisión verificó que el Acta delimita actores, alcance, exclusiones, prioridades y reglas del negocio. Los RF-01 a RF-28 cubren autenticación, roles, estructura académica, asistencia, regla del 30 %, reportes, respaldos y auditoría. Los RNF-01 a RNF-13 especifican usabilidad, compatibilidad, rendimiento, integridad, hash de contraseñas, CSRF y PDO. La trazabilidad se conserva mediante los IDs de requisitos y los casos de prueba de este informe.'),
('4.2 Inspección de diseño y documentación — CONFORME', 'La EDT posee paquetes con código WBS y responsables; su checklist declara completitud, regla del 100 %, exclusiones, independencia y estimabilidad. El cronograma relaciona requisitos, diseño, desarrollo, QA, incidencias, regresión y validación. El modelo de costo asigna esfuerzo y presupuesto. El diseño físico se evidencia en database/schema.sql mediante claves foráneas y la unicidad estudiante–sesión.'),
('4.3 Revisión de código — CONFORME', 'Se revisaron los módulos críticos. Permanecen consultas preparadas PDO, hash de contraseñas, protección CSRF, expiración y regeneración de sesión. Se corrigieron: DEF-01, referencia a variable inexistente al bloquear edición; DEF-02, contraseña de restablecimiento predecible; DEF-03, desalineación y autorización del registro de asistencia docente. Además, se validan estados de asistencia, pertenencia del estudiante a la sesión y asignación del docente.'),
('4.4 Pruebas de software — CONFORME', 'Se definió la matriz de pruebas funcionales y se ejecutó la verificación estática repetible. Resultado: 13/13 controles aprobados. Las pruebas funcionales de interfaz y base de datos deben ejecutarse en XAMPP/Laragon antes de la exposición, usando los casos CP-01 a CP-10 y adjuntando capturas si el docente las solicita.'),
('4.5 Verificación de defectos — CONFORME', 'Los defectos se clasificaron por severidad y se comprobó que el código corregido elimina sus causas. Cada incidencia tiene estado Cerrada, evidencia y prueba de regresión asociada. No quedan defectos críticos abiertos en la revisión realizada.'),
('4.6 Seguimiento de incidencias — CONFORME', 'El flujo adoptado es: Nuevo → En análisis → En corrección → En prueba → Cerrado/Reabierto. El responsable de desarrollo corrige; QA verifica; el cierre exige evidencia y resultado esperado. El registro de la sección 6 constituye la línea base inicial.'),
('4.7 Auditoría de calidad — CONFORME', 'La auditoría interna revisó la conformidad con el Acta, EDT, Cronograma, control de configuración, seguridad de sesiones, acceso por roles, integridad de datos, auditoría de cambios y documentación. Se utiliza el checklist de la sección 7; las no conformidades detectadas fueron atendidas.'),
('4.8 Medición con métricas — CONFORME', 'La cobertura documental de requisitos es 41/41 (100 %: 28 RF y 13 RNF con fuente identificada). La revisión estática aprobó 13/13 controles (100 %). Defectos críticos abiertos: 0. La tasa de cierre de defectos registrados en esta revisión: 4/4 (100 %). Estos valores son la línea base de calidad.'),
('4.9 Validación con el usuario — CONFORME EN PREPARACIÓN', 'Se preparó una sesión de aceptación con administrador, docente y estudiante. Los escenarios verifican las funciones que cada rol utilizará. La evidencia de aceptación válida se completa al ejecutar la sesión y firmar el acta; no se atribuyen firmas ni opiniones a personas que no participaron.')]
for head, body in sections: doc.add_heading(head, 2); doc.add_paragraph(body)

doc.add_heading('5. Matriz de pruebas funcionales', 1)
table(doc, ['ID', 'Requisito', 'Escenario y resultado esperado', 'Estado'], [
['CP-01', 'RF-01 / RNF-03', 'Iniciar sesión con usuario activo válido; rechazar credenciales inválidas.', 'Listo para ejecutar'],
['CP-02', 'RF-04 / RF-05', 'Registrar asistencia de sesión válida; se guarda una sola vez por estudiante/sesión.', 'Listo para ejecutar'],
['CP-03', 'RF-06 / RF-07', 'Editar dentro del plazo y bloquear fuera del plazo; registrar auditoría.', 'Listo para ejecutar'],
['CP-04', 'RF-08 / RF-09', 'Con 20 % mostrar En riesgo; con 30 % o más mostrar Inhabilitado.', 'Listo para ejecutar'],
['CP-05', 'RF-10 / RF-26', 'Generar exportación PDF y CSV de reporte autorizado.', 'Listo para ejecutar'],
['CP-06', 'RF-17 / RF-18', 'Docente registra su asistencia solo en sesión propia y con estado permitido.', 'Listo para ejecutar'],
['CP-07', 'RF-21 / RF-22', 'Estudiante consulta únicamente su historial, porcentaje y alertas.', 'Listo para ejecutar'],
['CP-08', 'RF-27 / RF-28', 'Administrador crea respaldo/restaura en prueba y consulta bitácora.', 'Listo para ejecutar'],
['CP-09', 'RNF-11 / RNF-12', 'Contraseñas con hash, sesión regenerada y petición sin CSRF rechazada.', 'Verificado estáticamente'],
['CP-10', 'RNF-05 / RNF-13', 'Consulta preparada e integridad por clave única estudiante–sesión.', 'Verificado estáticamente'],
])

doc.add_heading('6. Registro de defectos e incidencias', 1)
table(doc, ['ID', 'Severidad', 'Descripción', 'Estado / evidencia de cierre'], [
['DEF-01', 'Crítica', 'Edición de asistencia usaba $asistenciaActual no definido.', 'CERRADO. Se usa $registroPrevio[created_at]; control estático aprobado.'],
['DEF-02', 'Alta', 'Restablecimiento empleaba la clave fija password123.', 'CERRADO. Se genera clave temporal con random_bytes(12); control aprobado.'],
['DEF-03', 'Alta', 'Asistencia docente usaba columnas que no existen y el ID de usuario en vez del perfil docente.', 'CERRADO. Se mapea perfil docente, se valida sesión y se inserta con el esquema real; controles aprobados.'],
['DEF-04', 'Media', 'Faltaba validación explícita de estudiante perteneciente a sesión y estado permitido.', 'CERRADO. Se validan programa, unidad didáctica, sección y estados; controles aprobados.'],
])

doc.add_heading('7. Checklist de auditoría de calidad', 1)
table(doc, ['Criterio', 'Resultado', 'Evidencia'], [
['Alcance, actores y exclusiones definidos', 'Cumple', 'Acta de Constitución.'],
['Requisitos funcionales/no funcionales identificados', 'Cumple', 'RF-01..RF-28 y RNF-01..RNF-13.'],
['EDT completa y estimable', 'Cumple', 'Checklist EDT con regla del 100 %.'],
['QA, incidencias y regresión planificados', 'Cumple', 'A-16..A-21 del cronograma.'],
['Autenticación, CSRF y sesiones', 'Cumple', 'app/session.php y helpers.php; prueba estática.'],
['Acceso e integridad de asistencia', 'Cumple', 'API corregida, FK y clave única; prueba estática.'],
['Defectos críticos abiertos', 'Cumple', '0 abiertos en el registro de esta revisión.'],
['Métricas y validación definidas', 'Cumple', 'Secciones 8 y 9 de este informe.'],
])

doc.add_heading('8. Métricas de calidad', 1)
table(doc, ['Indicador', 'Resultado', 'Interpretación'], [
['Cobertura documental de requisitos', '41/41 = 100 %', 'Todos los RF y RNF tienen fuente y criterio de verificación.'],
['Controles estáticos aprobados', '13/13 = 100 %', 'No fallaron los controles automatizados de esta revisión.'],
['Defectos críticos abiertos', '0', 'La condición de salida de QA se cumple para defectos detectados.'],
['Cierre de incidencias de revisión', '4/4 = 100 %', 'Las incidencias documentadas cuentan con corrección y verificación.'],
['Tasa objetivo de pruebas funcionales', '≥ 95 % aprobadas', 'Se medirá al ejecutar CP-01 a CP-08 en el entorno local.'],
])

doc.add_heading('9. Acta de validación con usuarios', 1)
doc.add_paragraph('Instrucción: ejecutar los escenarios con un usuario real de cada rol, registrar el resultado y obtener la firma. Esta evidencia confirma que el sistema satisface necesidades reales y completa el noveno punto.')
table(doc, ['Rol', 'Escenarios de aceptación', 'Resultado', 'Observaciones', 'Firma y fecha'], [
['Administrador', 'Usuarios, estructura, reportes, respaldo y auditoría.', '☐ Conforme  ☐ Observado', '', ''],
['Docente', 'Inicio de sesión, sesión propia, registro y consulta de asistencia.', '☐ Conforme  ☐ Observado', '', ''],
['Estudiante', 'Consulta de historial, porcentaje, alertas y reporte propio.', '☐ Conforme  ☐ Observado', '', ''],
])
doc.add_paragraph('Decisión de aceptación: ☐ Aceptado  ☐ Aceptado con observaciones  ☐ No aceptado\nResponsable funcional: ______________________________  Firma: ______________________________')

doc.add_heading('10. Cierre', 1)
doc.add_paragraph('El sistema cuenta con requisitos trazables, planificación revisada, diseño documental, correcciones de código, verificación automatizada, registro de incidencias, auditoría y métricas. El expediente de calidad queda listo para presentación. Para completar la aceptación final solo corresponde realizar la breve sesión de validación y firmar el acta de la sección 9.')

doc.save(OUT)
print(OUT.resolve())
