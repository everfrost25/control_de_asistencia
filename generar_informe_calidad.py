from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path('Informe_Aseguramiento_Calidad_Control_Asistencia.docx')

def shade(cell, color):
    tc_pr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color); tc_pr.append(shd)

def table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; shade(c,'1F4E78')
        for r in c.paragraphs[0].runs: r.font.bold=True; r.font.color.rgb=RGBColor(255,255,255)
    for row in rows:
        cells=t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text=str(v); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.TOP
    doc.add_paragraph()
    return t

def bullet(doc, text):
    doc.add_paragraph(text, style='List Bullet')

doc=Document()
sec=doc.sections[0]; sec.top_margin=Inches(.7); sec.bottom_margin=Inches(.7); sec.left_margin=Inches(.8); sec.right_margin=Inches(.8)
styles=doc.styles
styles['Normal'].font.name='Aptos'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'Aptos'); styles['Normal'].font.size=Pt(10)
for s in ['Title','Heading 1','Heading 2']:
    styles[s].font.name='Aptos Display'

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('INFORME DE ASEGURAMIENTO DE CALIDAD'); r.bold=True; r.font.size=Pt(19); r.font.color.rgb=RGBColor(31,78,120)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Sistema Web de Control de Asistencia'); r.bold=True; r.font.size=Pt(14)
p=doc.add_paragraph('Curso: Calidad de Software\nFecha de evaluación: 29 de julio de 2026\nEstado: evaluación inicial basada en el repositorio disponible'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

doc.add_heading('1. Objetivo y alcance',1)
doc.add_paragraph('Evaluar el proyecto frente a los nueve procesos solicitados por el curso de Calidad de Software y dejar una base documental para completarlos. Se revisó el repositorio del sistema (PHP, SQL y README). Los cuatro archivos Word de gestión mencionados por el equipo no estuvieron disponibles en el entorno de revisión; por ello, las conclusiones sobre requisitos, diseño, cronograma y costos son provisionales y se identifican como “sin evidencia verificable”.')

doc.add_heading('2. Resultado ejecutivo',1)
table(doc,['Proceso solicitado','Resultado','Conclusión'],[
['1. Revisión de requisitos','Parcial','El README define funciones generales; falta especificación verificable de requisitos y criterios de aceptación.'],
['2. Inspección de diseño y documentación','Parcial','Hay esquema SQL y estructura de código; faltan diagramas/versionado/trazabilidad revisables.'],
['3. Revisión de código','No conforme','Se identificó un defecto funcional crítico y una práctica insegura de restablecimiento de contraseña.'],
['4. Pruebas de software','No conforme','No se encontraron casos, plan, evidencias ni pruebas automatizadas.'],
['5. Verificación de defectos','Parcial','Se detectaron defectos en esta revisión, pero no existe un registro formal de severidad, estado y evidencia.'],
['6. Seguimiento de incidencias','No conforme','No hay herramienta/bitácora de incidencias, responsable, fechas ni criterio de cierre.'],
['7. Auditoría de calidad','Parcial','Existe auditoría funcional de cambios de asistencia; no una auditoría del proceso de calidad contra una norma/lista.'],
['8. Medición con métricas','Parcial','El sistema muestra indicadores operativos; faltan métricas de calidad del producto y proceso.'],
['9. Validación con el usuario','Sin evidencia','No se encontró acta de aceptación, resultados UAT ni retroalimentación de usuarios.'],
])
doc.add_paragraph('Dictamen: el sistema implementa varias funciones del negocio y controles técnicos básicos, pero no demuestra todavía el cumplimiento integral de los nueve procesos. Los puntos 3, 4 y 6 requieren atención prioritaria antes de declarar el proyecto validado.', style='Intense Quote')

doc.add_heading('3. Evidencia positiva encontrada',1)
for x in [
'El README describe roles (administrador, docente y estudiante), control de asistencia, cálculo de inasistencias, importación y reportes.',
'La base de datos define integridad referencial y una restricción única por estudiante y sesión para evitar registros duplicados.',
'El sistema usa consultas preparadas con PDO, hash de contraseñas y control de sesión.',
'Las API protegidas revisadas requieren autenticación y token CSRF.',
'Existe una tabla de auditoría de cambios de asistencia y una pantalla para consultar la bitácora.',
'Se manejan parámetros configurables para reglas de inasistencia y edición.'
]: bullet(doc,x)

doc.add_heading('4. Hallazgos y defectos',1)
table(doc,['ID','Severidad','Hallazgo / evidencia','Impacto y acción requerida'],[
['DEF-01','Crítica','En public/api/asistencia.php se consulta $asistenciaActual[\'fecha_registro\'], pero esa variable no está definida; la consulta previa devuelve created_at.','Al editar una asistencia se produce error y se revierte la transacción. Corregir para usar $registroPrevio[\'created_at\']; crear prueba de edición dentro/fuera del plazo.'],
['DEF-02','Alta','public/api/usuarios.php restablece cualquier contraseña a la constante “password123” y la devuelve en la respuesta.','Compromete cuentas. Generar una clave aleatoria temporal o enlace de recuperación; obligar cambio al inicio; no devolver la contraseña en texto.'],
['DEF-03','Media','La asistencia valida que la sesión exista, pero no se observa en ese endpoint una validación de que el docente sea titular de esa sesión ni de que cada estudiante pertenezca a ella.','Podría registrarse asistencia sobre registros no autorizados o inconsistentes. Validar relación docente–sesión–estudiante antes de insertar/actualizar.'],
['DEF-04','Media','El README expone credenciales de demostración simples y el esquema incluye datos de prueba.','Riesgo si se despliega fuera de un ambiente académico. Separar semillas de demo y exigir cambio de credenciales.'],
['DEF-05','Media','No se encontraron pruebas unitarias, de integración, UAT, registro de defectos ni historial de incidencias.','No se puede demostrar cobertura, repetibilidad ni corrección. Implementar los artefactos de las secciones 5–8.'],
['DEF-06','Baja','El error de conexión a base de datos devuelve detalle técnico al usuario.','Expone información de infraestructura. Registrar detalle internamente y mostrar un mensaje genérico.'],
])

doc.add_heading('5. Aplicación de los nueve puntos',1)
items=[
('5.1 Revisión de requisitos','Estado: parcial. El README permite identificar necesidades funcionales generales, pero no sustituye una especificación. Falta una matriz con identificador, descripción, prioridad, fuente, regla de negocio, criterio de aceptación y estado de aprobación. Acción: transformar el acta y la EDT en requisitos trazables; por cada requisito definir escenarios “Dado/Cuando/Entonces”.'),
('5.2 Inspección de diseño y documentación','Estado: parcial. El esquema SQL aporta diseño de datos y la estructura de carpetas separa configuración, repositorios, API y vistas. Faltan diagrama de arquitectura, modelo entidad-relación legible, diagramas de casos de uso/flujo y control de versiones de documentos. Acción: inspección por pares usando checklist y aprobación con fecha/firma.'),
('5.3 Revisión de código','Estado: no conforme. La revisión estática identificó DEF-01 a DEF-04. Acción: revisión por pull request o checklist; bloquear cambios sin revisión y sin prueba asociada. Criterios mínimos: validación de autorización, consultas parametrizadas, manejo seguro de secretos, errores no expuestos y nombres coherentes.'),
('5.4 Pruebas de software','Estado: no conforme. No existe evidencia de ejecución de pruebas. Acción: elaborar plan y registrar resultado de cada caso. Ejecutar regresión después de corregir DEF-01 y DEF-02.'),
('5.5 Verificación de defectos','Estado: parcial. Este informe inicia el registro, pero falta confirmar cada corrección en una versión definida. Acción: para cada defecto registrar pasos para reproducir, resultado esperado/actual, severidad, evidencia, responsable y verificador independiente.'),
('5.6 Seguimiento de incidencias','Estado: no conforme. Acción: usar una hoja compartida o tablero con flujo Nuevo → En análisis → En corrección → En prueba → Cerrado/Reabierto; cada cierre requiere evidencia de prueba y responsable de QA.'),
('5.7 Auditoría de calidad','Estado: parcial. La bitácora funcional no equivale a auditoría de proceso. Acción: auditoría interna con checklist de requisitos, diseño, código, pruebas, seguridad y documentación; emitir no conformidades y plan de acción.'),
('5.8 Medición con métricas','Estado: parcial. Hay indicadores académicos, pero no métricas de calidad. Acción: medir cobertura de requisitos, tasa de aprobación, densidad de defectos, defectos reabiertos y tiempo de resolución. No inventar valores: registrar una línea base tras ejecutar pruebas.'),
('5.9 Validación con el usuario','Estado: sin evidencia. Acción: realizar prueba de aceptación con al menos un administrador, docente y estudiante; documentar escenarios, resultado, observaciones y acta de aceptación o lista de pendientes.')]
for h,t in items: doc.add_heading(h,2); doc.add_paragraph(t)

doc.add_heading('6. Plan mínimo de pruebas',1)
table(doc,['ID','Escenario','Resultado esperado'],[
['CP-01','Inicio de sesión con credenciales válidas e inválidas','Permite acceso solo al usuario activo válido; mensaje genérico ante error.'],
['CP-02','Registro de asistencia de una sesión válida','Se registra una sola asistencia por estudiante/sesión y la sesión cambia a Registrada.'],
['CP-03','Edición dentro del tiempo configurado','Actualiza asistencia y guarda auditoría; no genera error.'],
['CP-04','Edición después del tiempo configurado','Bloquea el cambio e informa el motivo sin alterar datos.'],
['CP-05','Docente intenta modificar sesión ajena','El servidor rechaza la acción con 403/validación de asignación.'],
['CP-06','Cálculo 20% y 30% de inasistencias','Clasifica respectivamente En riesgo e Inhabilitado según configuración.'],
['CP-07','Restablecimiento de contraseña','No usa clave fija, obliga cambio posterior y no expone la nueva clave.'],
['CP-08','CSRF y rol no autorizado contra APIs','Solicitud sin token o rol sin permiso es rechazada.'],
])

doc.add_heading('7. Registro y seguimiento de incidencias (plantilla)',1)
table(doc,['Campo','Contenido requerido'],[
['Identificador','Ejemplo: DEF-007'],['Título y módulo','Descripción breve y ruta/pantalla afectada'],['Severidad','Crítica, Alta, Media o Baja'],['Estado','Nuevo, En análisis, En corrección, En prueba, Cerrado o Reabierto'],['Evidencia','Captura, solicitud, datos de prueba o enlace a commit'],['Responsables','Reportante, desarrollador y verificador'],['Fechas','Reporte, compromiso, corrección y cierre'],['Criterio de cierre','Prueba repetida con resultado esperado y sin regresión']],)

doc.add_heading('8. Métricas propuestas',1)
table(doc,['Métrica','Fórmula','Meta inicial sugerida'],[
['Cobertura de requisitos','Requisitos con al menos una prueba ejecutada / requisitos totales × 100','100% para requisitos críticos; ≥90% total'],
['Tasa de aprobación','Casos aprobados / casos ejecutados × 100','≥95% antes de entrega'],
['Densidad de defectos','Defectos confirmados / tamaño del producto (KLOC o módulos)','Línea base primero; tendencia descendente'],
['Defectos reabiertos','Defectos reabiertos / defectos cerrados × 100','≤10%'],
['Tiempo medio de resolución','Suma de horas desde reporte hasta cierre / defectos cerrados','Definir SLA: críticos ≤24 h en entorno académico'],
])

doc.add_heading('9. Validación con usuarios: acta breve',1)
doc.add_paragraph('Producto: Sistema Web de Control de Asistencia\nVersión evaluada: __________    Fecha: __________\nParticipantes: Administrador __________ / Docente __________ / Estudiante __________')
table(doc,['Rol','Escenarios ejecutados','Resultado','Observaciones / pendientes','Firma'],[['Administrador','','','',''],['Docente','','','',''],['Estudiante','','','','']])
doc.add_paragraph('Conclusión de aceptación: ☐ Aceptado  ☐ Aceptado con observaciones  ☐ No aceptado\nResponsable funcional: ____________________    Firma: ____________________')

doc.add_heading('10. Prioridad de mejora',1)
doc.add_paragraph('1) Corregir DEF-01 y DEF-02.  2) Crear y ejecutar CP-01 a CP-08.  3) Llevar el registro de defectos hasta cierre.  4) Completar matriz de requisitos y anexar la documentación original.  5) Obtener el acta de validación de usuarios. Con estas evidencias el equipo podrá demostrar los nueve puntos de forma defendible ante el curso.')

doc.save(OUT)
print(OUT.resolve())
